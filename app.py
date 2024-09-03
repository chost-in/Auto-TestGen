import streamlit as st
import json
import os
import numpy as np
from docx import Document
import pdfplumber
import matplotlib.colors as mcolors
from operator import index
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import string
from dotenv import load_dotenv
import docx
from PIL import Image
from openai import OpenAI
import base64

load_dotenv()


session_state = st.session_state
if "user_index" not in st.session_state:
    st.session_state["user_index"] = 0


def signup(json_file_path="data.json"):
    st.title("Signup Page")
    with st.form("signup_form"):
        st.write("Fill in the details below to create an account:")
        name = st.text_input("Name:")
        email = st.text_input("Email:")
        age = st.number_input("Age:", min_value=0, max_value=120)
        sex = st.radio("Sex:", ("Male", "Female", "Other"))
        password = st.text_input("Password:", type="password")
        confirm_password = st.text_input("Confirm Password:", type="password")

        if st.form_submit_button("Signup"):
            if password == confirm_password:
                user = create_account(name, email, age, sex, password, json_file_path)
                session_state["logged_in"] = True
                session_state["user_info"] = user
            else:
                st.error("Passwords do not match. Please try again.")


def check_login(username, password, json_file_path="data.json"):
    try:
        with open(json_file_path, "r") as json_file:
            data = json.load(json_file)

        for user in data["users"]:
            if user["email"] == username and user["password"] == password:
                session_state["logged_in"] = True
                session_state["user_info"] = user
                st.success("Login successful!")
                return user

        st.error("Invalid credentials. Please try again.")
        return None
    except Exception as e:
        st.error(f"Error checking login: {e}")
        return None


def initialize_database(json_file_path="data.json"):
    try:
        # Check if JSON file exists
        if not os.path.exists(json_file_path):
            # Create an empty JSON structure
            data = {"users": []}
            with open(json_file_path, "w") as json_file:
                json.dump(data, json_file)
    except Exception as e:
        print(f"Error initializing database: {e}")


def create_account(name, email, age, sex, password, json_file_path="data.json"):
    try:
        # Check if the JSON file exists or is empty
        if not os.path.exists(json_file_path) or os.stat(json_file_path).st_size == 0:
            data = {"users": []}
        else:
            with open(json_file_path, "r") as json_file:
                data = json.load(json_file)

        # Append new user data to the JSON structure
        user_info = {
            "name": name,
            "email": email,
            "age": age,
            "sex": sex,
            "password": password,
            "test_cases": None,
            "program": None,
        }
        data["users"].append(user_info)

        # Save the updated data to JSON
        with open(json_file_path, "w") as json_file:
            json.dump(data, json_file, indent=4)

        st.success("Account created successfully! You can now login.")
        return user_info
    except json.JSONDecodeError as e:
        st.error(f"Error decoding JSON: {e}")
        return None
    except Exception as e:
        st.error(f"Error creating account: {e}")
        return None


def login(json_file_path="data.json"):
    st.title("Login Page")
    username = st.text_input("Username:")
    password = st.text_input("Password:", type="password")

    login_button = st.button("Login")

    if login_button:
        user = check_login(username, password, json_file_path)
        if user is not None:
            session_state["logged_in"] = True
            session_state["user_info"] = user
        else:
            st.error("Invalid credentials. Please try again.")


def get_user_info(email, json_file_path="data.json"):
    try:
        with open(json_file_path, "r") as json_file:
            data = json.load(json_file)
            for user in data["users"]:
                if user["email"] == email:
                    return user
        return None
    except Exception as e:
        st.error(f"Error getting user information: {e}")
        return None


def render_dashboard(user_info, json_file_path="data.json"):
    try:
        st.title(f"Welcome to the Dashboard, {user_info['name']}!")
        st.subheader("User Information:")
        st.write(f"Name: {user_info['name']}")
        st.write(f"Sex: {user_info['sex']}")
        st.write(f"Age: {user_info['age']}")
    except Exception as e:
        st.error(f"Error rendering dashboard: {e}")


def extract_text(file) -> str:
    if isinstance(file, str):
        file_extension = os.path.splitext(file)[1].lower()
    else:
        file_extension = os.path.splitext(file.name)[1].lower()

    if file_extension == ".pdf":
        if isinstance(file, str):
            with pdfplumber.open(file) as pdf:
                text = "\n".join(
                    page.extract_text() for page in pdf.pages if page.extract_text()
                )
        else:
            with pdfplumber.open(file) as pdf:
                text = "\n".join(
                    page.extract_text() for page in pdf.pages if page.extract_text()
                )
    elif file_extension == ".docx":
        if isinstance(file, str):
            doc = docx.Document(file)
        else:
            doc = docx.Document(file)
        text = "\n".join([para.text for para in doc.paragraphs])
    elif file_extension == ".txt":
        if isinstance(file, str):
            with open(file, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
        else:
            with file as f:
                text = f.read()
    else:
        if isinstance(file, str):
            with open(file, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
        else:
            with file as f:
                text = f.read()
    return text


def get_unit_test_cases(program, language):
    try:
        client = OpenAI(
            api_key=os.environ.get("OPENAI_API_KEY"),
        )
        prompt = f"Given the following {language} program:\n{program}\nGenerate all possible unit test cases for this program. Make sure to cover all diverse scenarios. The test cases should be in the form of input and expected output."
        messages = [{"role": "system", "content": prompt}]
        response = client.chat.completions.create(
            messages=messages,
            model="gpt-3.5-turbo-0125",
        )
        return response.choices[0].message.content
    except Exception as e:
        # Handle exceptions such as API errors or connection issues
        print(f"An error occurred: {e}")
        return None


def main(json_file_path="data.json"):
    st.sidebar.title("Unit Test Cases Generator")
    page = st.sidebar.radio(
        "Go to",
        ("Signup/Login", "Dashboard", "Upload a program", "View Test Cases"),
        key="Unit Test Cases Generator",
    )

    if page == "Signup/Login":
        st.title("Signup/Login Page")
        login_or_signup = st.radio(
            "Select an option", ("Login", "Signup"), key="login_signup"
        )
        if login_or_signup == "Login":
            login(json_file_path)
        else:
            signup(json_file_path)

    elif page == "Dashboard":
        if session_state.get("logged_in"):
            render_dashboard(session_state["user_info"], json_file_path)
        else:
            st.warning("Please login/signup to view the dashboard.")

    elif page == "Upload a program":
        if session_state.get("logged_in"):
            st.title("Upload a program")
            uploaded_file = st.file_uploader("Choose a file", type=None)
            language = st.selectbox(
                "Select the language of the program",
                ("Select a language", "Python", "Java", "C++", "C#", "Javascript"),
            )
            if uploaded_file is not None and language != "Select a language":
                program = extract_text(uploaded_file)
                st.code(program, language=language.lower())
                st.write("File name: ", uploaded_file.name)
                st.image(Image.open("Images/logo.png"), use_column_width=True)
                st.success("Program uploaded successfully!")
                st.title("Unit Test Cases")
                st.subheader("The following are the unit test cases for the program:")
                unit_test_cases = get_unit_test_cases(program, language)
                with open(json_file_path, "r+") as json_file:
                    data = json.load(json_file)
                    user_index = next(
                        (
                            i
                            for i, user in enumerate(data["users"])
                            if user["email"] == session_state["user_info"]["email"]
                        ),
                        None,
                    )
                    if user_index is not None:
                        user_info = data["users"][user_index]
                        if user_info["program"] is None:
                            user_info["program"] = []
                        user_info["program"].append(str(program))

                        if "test_cases" not in user_info or user_info["test_cases"] is None:
                            user_info["test_cases"] = []
                        user_info["test_cases"].append(unit_test_cases)

                        session_state["user_info"] = user_info

                        json_file.seek(0)
                        json_file.truncate()  # Truncate the file before writing
                        json.dump(data, json_file, indent=4)
                    else:
                        st.error("User not found.")

                st.code(unit_test_cases, language=language.lower())
        else:
            st.warning("Please login/signup to view the dashboard.")

    elif page == "View Test Cases":
        if session_state.get("logged_in"):
            user_info = get_user_info(
                session_state["user_info"]["email"], json_file_path
            )
            st.title("View Test Cases")
            if user_info is not None:
                programs = user_info["program"]
                test_cases = user_info["test_cases"]
                if programs is not None:
                    st.subheader("Programs")
                    for i in range(len(programs)):
                        try:
                            st.subheader(f"Program {i + 1}")
                            st.code(programs[i])
                            st.write("\n")
                            st.subheader("Unit Test Cases")
                            st.write(test_cases[i])
                            st.write("\n")
                        except Exception as e:
                            pass
                else:
                    st.warning("No programs uploaded yet.")
            else:
                st.error("User not found.")


if __name__ == "__main__":
    initialize_database()
    main()
